import streamlit as st
from crewai import Agent, Task, Crew, Process
from langchain_community.tools.google_trends import GoogleTrendsQueryRun
from langchain_community.utilities.google_trends import GoogleTrendsAPIWrapper
from langchain_openai import ChatOpenAI
from crewai_tools import SerperDevTool, ScrapeWebsiteTool
from tools.semrush_keyword import SemrushKeyWordTools
from tools.semrush_tools import SemrushTools
import re
import sys
import os
from dotenv import load_dotenv
from langchain_anthropic import ChatAnthropic
import docx
import requests
import csv

from datetime import datetime

# Ensure the results directory exists
if not os.path.exists("Results"):
    os.makedirs("Results")
    
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

load_dotenv()

def setup_llm(llm_option):
    if llm_option == 'OpenAI GPT-4o':
        api_key = os.getenv('OPENAI_API_KEY')
        return ChatOpenAI(model="gpt-4o", api_key=api_key)
    else:
        api_key = os.getenv('CLAUDE_API_KEY')
        return ChatAnthropic(model="claude-3-haiku-20240307", api_key=api_key)

def fetch_related_keywords(api_key, phrase, lang):
    url = f"https://api.semrush.com/?type=phrase_related&key={api_key}&phrase={phrase}&export_columns=Ph,Nq,Nr,Td,Rr&database={lang}&display_limit=10&display_sort=nq_desc&display_filter=%2B|Nq|Lt|1000"
    response = requests.get(url)
    try:
        response.raise_for_status()
        csv_content = response.text.splitlines()
        reader = csv.DictReader(csv_content, delimiter=';')
        return [row for row in reader]
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error occurred: {http_err}")
    except requests.exceptions.RequestException as req_err:
        st.error(f"Request error occurred: {req_err}")
    except Exception as e:
        st.error(f"Error parsing CSV response for related keywords: {e}")
        st.write("Raw response:", response.text)
    return []

# Function to fetch and parse QA from SEMrush API
def fetch_qa(api_key, phrase, lang):
    url = f"https://api.semrush.com/?type=phrase_questions&key={api_key}&phrase={phrase}&export_columns=Ph,Nq,Nr,Td&database={lang}&display_limit=10&display_sort=nq_desc&display_filter=%2B|Nq|Lt|1000"
    response = requests.get(url)
    try:
        response.raise_for_status()
        csv_content = response.text.splitlines()
        reader = csv.DictReader(csv_content, delimiter=';')
        return [row for row in reader]
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error occurred: {http_err}")
    except requests.exceptions.RequestException as req_err:
        st.error(f"Request error occurred: {req_err}")
    except Exception as e:
        st.error(f"Error parsing CSV response for QA data: {e}")
        st.write("Raw response:", response.text)
    return []


st.title("SEO Briefing Generator")

# Language selection
language = st.selectbox("Select Language / Sprache wählen", ["English", "Deutsch"])
is_german = language == "Deutsch"

llm_options = ['OpenAI GPT-4o', 'Claude-3', 'Groq']
llm_choice = st.selectbox("Select the LLM to use:", llm_options)

llm = setup_llm(llm_choice)

google_search = SerperDevTool()
website_scrapper = ScrapeWebsiteTool()
google_trends_api_wrapper = GoogleTrendsAPIWrapper()
google_trends_tool = GoogleTrendsQueryRun(api_wrapper=google_trends_api_wrapper)


class StreamToExpander:
    def __init__(self, expander, buffer_limit=10000):
        self.expander = expander
        self.buffer = []
        self.buffer_limit = buffer_limit

    def write(self, data):
        cleaned_data = re.sub(r'\x1B\[\d+;?\d*m', '', data)
        if len(self.buffer) >= self.buffer_limit:
            self.buffer.pop(0)
        self.buffer.append(cleaned_data)

        if "\n" in data:
            self.expander.markdown(''.join(self.buffer), unsafe_allow_html=True)
            self.buffer.clear()

    def flush(self):
        if self.buffer:
            self.expander.markdown(''.join(self.buffer), unsafe_allow_html=True)
            self.buffer.clear()

st.markdown("""
    <style>
        .stApp {
            background-color: white;
            color: black;
        }
        .stButton button {
            background-color: #003781;
            color: white;
        }
        .stSelectbox div, .stTextInput div, .stTextArea div, .stSlider div {
            border-color: black !important;
        }
    </style>
    """, unsafe_allow_html=True)

with st.form("research_form"):
    focus_keyword_label = "Enter a focus keyword:" if not is_german else "Geben Sie ein Fokus-Keyword ein:"
    target_audience_label = "Describe your target audience:" if not is_german else "Beschreiben Sie Ihre Zielgruppe:"
    tone_label = "Select the tone of the blog post:" if not is_german else "Wählen Sie den Ton des Blogposts:"
    length_label = "Select the desired length of the blog post (in words):" if not is_german else "Wählen Sie die gewünschte Länge des Blogposts (in Wörtern):"
    key_points_label = "Enter key points or topics you want to cover in the blog post:" if not is_german else "Geben Sie wichtige Punkte oder Themen ein, die Sie im Blogpost behandeln möchten:"
    submit_button_label = "Generate SEO Briefing" if not is_german else "SEO-Briefing generieren"
    brand_name_label = "Enter Brand Name:" if not is_german else "Geben Sie den Markennamen ein:"

    focus_keyword = st.text_input(focus_keyword_label, "Renting trailers insurance")
    target_audience = st.text_input(target_audience_label, "Small business owners")
    tone = st.selectbox(tone_label, ["Formal", "Informal", "Professional", "Conversational"])
    length = st.slider(length_label, min_value=300, max_value=3000, value=1000, step=100)
    key_points = st.text_area(key_points_label, "Benefits of renting trailers, insurance options, cost considerations, tips for renting")
    brand_name = st.text_input(brand_name_label, "Your Brand Name")
    submit_button = st.form_submit_button(submit_button_label)

if submit_button:
    process_output_expander = st.expander("Processing Output:")
    sys.stdout = StreamToExpander(process_output_expander)
    
    try:
        semrush_api_key = os.getenv('SEMRUSH_API_KEY')
        lang = 'de' if is_german else 'us'
        related_keywords = fetch_related_keywords(semrush_api_key, focus_keyword, lang)
        qa_data = fetch_qa(semrush_api_key, focus_keyword, lang)
        
        
        boss_agent = Agent(
            role="Boss Agent",
            goal="Lead the development of an effective SEO strategy to improve website visibility and search engine ranking. This includes overseeing the content creation process, setting deadlines, and ensuring quality standards are met. The goal is to produce high-quality content that meets client requirements and surpasses expectations, ultimately driving organic traffic and conversions.",
            backstory="The SEO Strategy Manager is an experienced SEO professional with expertise in devising and executing successful SEO strategies. They possess a deep understanding of SEO best practices and stay updated on the latest trends and algorithm changes. Their leadership skills enable them to manage teams effectively and adapt strategies to achieve optimal results.",
            tools=[
                google_search,
                website_scrapper,
                google_trends_tool,
            ],
            llm=llm,
            verbose=True,
            allow_delegation=False,
        )

        researcher_agent = Agent(
            role="Researcher Agent",
            goal="Conduct in-depth research on relevant topics to inform the SEO strategy. This involves gathering data, identifying key trends, and analyzing competitor strategies. The goal is to provide valuable insights that contribute to the development of an effective SEO plan.",
            backstory="The SEO Research Analyst is a skilled researcher with a knack for uncovering valuable insights from data. They possess strong analytical skills and a keen eye for detail, allowing them to identify emerging trends and opportunities. Their research expertise is instrumental in shaping the SEO strategy and driving website performance improvements.",
            tools=[
                google_search,
                website_scrapper,
                google_trends_tool,
            ],
            llm=llm,
            verbose=True,
            allow_delegation=False,
        )

        technical_seo_agent = Agent(
            role="Technical SEO Agent",
            goal="Generate a comprehensive SEO brief report focused on optimizing website content for search engines. This includes analyzing meta tags, descriptions, related keywords, and search engine trends. The goal is to provide actionable recommendations to enhance website visibility and ranking in search engine results.",
            backstory="The SEO Technical Analyst is an expert in technical SEO with a deep understanding of search engine algorithms and ranking factors. They possess advanced analytical skills and leverage data-driven insights to optimize website performance. Their expertise in identifying relevant keywords, optimizing meta tags, and leveraging search engine trends is crucial in improving website visibility and driving organic traffic.",
            tools=[
                google_search,
                website_scrapper,
                google_trends_tool,
            ],
            llm=llm,
            verbose=True,
            allow_delegation=False,
        )

        outliner_agent = Agent(
            role="Outliner Agent",
            goal=f"Create the initial outline for the SEO optimized landing page. This includes researching the topic, identifying key points, and structuring the content in a logical and engaging way. The landing page should be tailored to {target_audience} and have a {tone} tone.",
            backstory="The Outliner Agent is a skilled writer with a talent for structuring content in a logical and engaging way. They have a deep understanding of the content creation process and are able to identify key points and themes. They are also knowledgeable about the latest trends and best practices in content creation and are able to adapt to changing circumstances.",
            tools=[google_search, website_scrapper, google_trends_tool],
            llm=llm,
            verbose=True,
            allow_delegation=False,
        )

        outline_task = Task(
            description=f"Create the initial outline for the blog post. This includes researching the topic, identifying key points, and structuring the content in a logical and engaging way. The blog post should be tailored to {target_audience} and have a {tone} tone. Focus on the key points: {key_points}.",
            expected_output="A detailed outline for the blog post, including the main points and subpoints, as well as any relevant research or data.",
            agent=outliner_agent,
            output_file=f"Results/outline-[{timestamp}].md"
        )

        keyword_research_task = Task(
            description=f"Conduct thorough keyword research to identify relevant keywords for the landing page focused on {focus_keyword}. This includes analyzing search volume, competition, and relevance to the topic. The landing page should be tailored to {target_audience} and have a {tone} tone. Focus on the key points: {key_points}.",
            expected_output="A list of relevant keywords, along with their search volume and competition metrics.",
            agent=technical_seo_agent,
            output_file=f"Results/keyword_research-[{timestamp}].md"
        )

        technical_seo_task = Task(
            description=f"Ensure that the blog post is optimized for search engines. This includes identifying relevant keywords, optimizing the meta tags and descriptions, and ensuring that the content is structured in a way that is easy for search engines to crawl and index. The blog post should be tailored to {target_audience} and have a {tone} tone. Focus on the key points: {key_points}. Semrush Fethched Keywords: {related_keywords} and Semrush QA_data: {qa_data}",
            expected_output="""
                    - Meta Title
                    - Meta Description
                    - Results of Competitor Search
                    - Related Keywords (Proof Keywords)   // via semrush api features: related keywords
                    - Headline hierarchies // via LLM suggestions
                    - QA // via semrush api features:QA """,
            agent=technical_seo_agent, 
            output_file=f"Results/technical_seo-[{timestamp}].md"
        )

        # Define the Crew
        research_crew = Crew(
            agents=[
                boss_agent,
                outliner_agent,
                researcher_agent,
                technical_seo_agent,
            ],
            tasks=[
                outline_task,
                keyword_research_task,
            ],
            process=Process.hierarchical,
            manager_llm=llm,
        )

        result = research_crew.kickoff()
        st.write(result)
        # Create a document with the required sections
        doc = docx.Document()
        doc.add_heading('SEO Briefing', 0)
        doc.add_heading('Meta Title', level=1)
        doc.add_paragraph(f'1er BMW Versicherung und Kosten | {brand_name}')
        doc.add_heading('Meta Description', level=1)
        doc.add_paragraph(f'Optimize your landing page for {focus_keyword} and attract {target_audience}.')
        doc.add_heading('Results of Competitor Search', level=1)
        # Include the result directly in the document
        doc.add_paragraph(result)

        doc.add_heading('Related Keywords (Proof Keywords)', level=1)
        doc.add_paragraph(f'Related keywords for {focus_keyword}:')
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Keyword'
        hdr_cells[1].text = 'Search Volume'
        hdr_cells[2].text = 'Number of Results'
        hdr_cells[3].text = 'Trend'
        hdr_cells[4].text = 'Relevance'
        for item in related_keywords:
            row_cells = table.add_row().cells
            row_cells[0].text = item['Ph']
            row_cells[1].text = str(item['Nq'])
            row_cells[2].text = str(item['Nr'])
            row_cells[3].text = str(item['Td'])
            row_cells[4].text = str(item['Rr'])

        doc.add_heading('Headline Hierarchies', level=1)
        # Include the headline hierarchies generated by LLM here
        doc.add_paragraph('Headline hierarchies generated by the LLM will be placed here.')

        doc.add_heading('QA', level=1)
        doc.add_paragraph(f'Questions and Answers related to {focus_keyword}:')
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Search Volume'
        hdr_cells[2].text = 'Number of Results'
        hdr_cells[3].text = 'Trend'
        for item in qa_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['Ph']
            row_cells[1].text = str(item['Nq'])
            row_cells[2].text = str(item['Nr'])
            row_cells[3].text = str(item['Td'])

        # Save the document
        doc_file = f"Results/SEO_Briefing_{timestamp}.docx"
        doc.save(doc_file)

        
        st.success(f"SEO Briefing has been generated successfully! [Download the document](/{doc_file})")

    except Exception as e:
        st.error(f"An error occurred: {e}")

